import os
import json
from docx import Document
from openai import OpenAI

# Configuration
# MODEL_NAME = "gemma2"
# MODEL_NAME = "mistral"
MODEL_NAME = "llama3.1"
BASE_URL = "http://10.211.59.1:11434/v1"
API_KEY = "ollama"
base_file_op_path = "./"

client = OpenAI(
    base_url=BASE_URL,
    api_key=API_KEY,  # required, but unused
)

# System prompts
SYSTEM_PROMPT_ACTION_ITEMS = """
Identify and summarize any action items from the meeting transcription.
List the tasks or action items discussed in the meeting.
Extract actionable items or to-do tasks from the meeting.
"""

SYSTEM_PROMPT_SUMMARY_NOTES ="""
Generate a summary of the key points discussed in the meeting.
Summarize the main topics covered in the meeting transcription.
Provide a brief overview of the meeting based on the transcription.
"""

SYSTEM_PROMPT_IDENTITY_PARTICIPANTS = """
List the participants or speakers in the meeting and their contributions.
Extract information about who spoke during the meeting and what they discussed.
"""

SYSTEM_PROMPT_DECISIONS = """
Identify and summarize any decisions made during the meeting.
Extract details about the decisions taken in the meeting.
Highlight key decision points from the meeting transcription.
"""

SYSTEM_PROMPT_IMPORTANT_NOTE = """
Extract timestamps and associated content from the meeting transcription.
Identify specific timings mentioned during the meeting and their context.
"""

SYSTEM_PROMPT_FOLLOW_UPS = """
List any follow-up tasks or items mentioned at the end of the meeting.
Extract information about follow-up actions discussed in the meeting.
"""

SYSTEM_PROMPT_EXTRACT_QUESTIONS = """
Identify any questions raised during the meeting and provide their context.
List the queries or questions discussed in the meeting transcription.
"""

SYSTEM_PROMPT_CAPTURE_AGENDA = """
Summarize the agenda items covered in the meeting transcription.
Extract information about the topics or agenda items discussed during the meeting.
"""

SYSTEM_PROMPT_HIGHLIGHT_CONCERNS = """
Identify and summarize any concerns or issues raised during the meeting.
Extract details about concerns voiced by participants in the meeting.
"""

SYSTEM_PROMPT_MEETING_SUMMARY = """
Generate a concise summary of the entire meeting based on the transcription.
Summarize the key takeaways and highlights from the meeting.
"""

SYSTEM_PROMPT_MOM = """
Generate detailed minutes of the meeting suitable for sharing with the client. Include key discussion points, decisions made, action items assigned to participants, and any notable follow-up tasks. Provide a comprehensive summary that captures the essence of the meeting and is clear and concise for client communication.
"""

def read_file(file_path):
    try:
        with open(file_path, 'r') as file:
            file_content = file.read()
        print("File read succesfully")
        return file_content
    except FileNotFoundError:
        print(f"The file '{file_path}' does not exist.")
        return None
    except Exception as e:
        print(f"An error occurred: {e}")
        return None  

def get_contents(transcription, system_prompt):
    response = client.chat.completions.create(
        model=MODEL_NAME,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        doc.add_paragraph()
    doc.save(filename)
    print("File written succesfully")

def meeting_minutes(transcription):
    print("Using Ollama Model:",MODEL_NAME)
    ACTION_ITEMS = get_contents(transcription, SYSTEM_PROMPT_ACTION_ITEMS)
    print("Got action Items")
    SUMMARY_NOTES = get_contents(transcription, SYSTEM_PROMPT_SUMMARY_NOTES)
    print("Got Summary Notes")
    IDENTITY_PARTICIPANTS = get_contents(transcription, SYSTEM_PROMPT_IDENTITY_PARTICIPANTS)
    print("Got Participants")
    DECISIONS = get_contents(transcription, SYSTEM_PROMPT_DECISIONS)
    print("Got Decision Items")
    IMPORTANT_NOTE = get_contents(transcription, SYSTEM_PROMPT_IMPORTANT_NOTE)
    print("Got Important Notes")
    FOLLOW_UPS = get_contents(transcription, SYSTEM_PROMPT_FOLLOW_UPS)
    print("Got Follow_up Items")
    EXTRACT_QUESTIONS = get_contents(transcription, SYSTEM_PROMPT_EXTRACT_QUESTIONS)
    print("Got Questions")
    CAPTURE_AGENDA = get_contents(transcription, SYSTEM_PROMPT_CAPTURE_AGENDA)
    print("Got CAPTURE_AGENDA")
    HIGHLIGHT_CONCERNS = get_contents(transcription, SYSTEM_PROMPT_HIGHLIGHT_CONCERNS)
    print("Got HIGHLIGHT_CONCERNS Items")
    MEETING_SUMMARY = get_contents(transcription, SYSTEM_PROMPT_MEETING_SUMMARY)
    print("Got MEETING_SUMMARY")
    MOM = get_contents(transcription, SYSTEM_PROMPT_MOM)
    print("Got MOM")
   
    return {
        'Action Items': ACTION_ITEMS,
        'Summary Notes': SUMMARY_NOTES,
        'Participants': IDENTITY_PARTICIPANTS,
        'Decisions': DECISIONS,
        'Important Note': IMPORTANT_NOTE,
        'Follow Ups': FOLLOW_UPS,
        'Questions': EXTRACT_QUESTIONS,
        'Agenda': CAPTURE_AGENDA,
        'Concerns': HIGHLIGHT_CONCERNS,
        'Meeting Summary': MEETING_SUMMARY,
        'MOM': MOM
    }

file_path = "IQUW_ VeriskMonthlyReviewMeeting.txt"
content = read_file(file_path)

if content:
    file_name = file_path.split("/")[-1].split(".")[0].replace(" ", "")
    minutes = meeting_minutes(content)
    save_as_docx(minutes, f'{base_file_op_path}{file_name}_detailed_meeting_minutes.docx')
    print("Saving document")
    with open(f'{base_file_op_path}{file_name}_minutes.json', 'w') as file:
        json.dump(minutes, file, indent=4)
    print("Document Saved")